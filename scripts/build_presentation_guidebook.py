from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    ListFlowable,
    ListItem,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUT_DOCX = ROOT / "output" / "documents" / "rent-car-guidebook-presentasi.docx"
OUT_PDF = ROOT / "output" / "pdf" / "rent-car-guidebook-presentasi.pdf"
OUT_MD = DOCS_DIR / "GUIDE_BOOK_PRESENTASI.md"

TITLE = "Guide Book Presentasi Rent Car Platform"
SUBTITLE = "Panduan naratif untuk mahasiswa, demo aplikasi, dan tanya jawab"
VERSION = "Versi 1.2 - 12 Mei 2026"


def build_content() -> list[dict]:
    return [
        {
            "type": "front",
            "title": TITLE,
            "subtitle": SUBTITLE,
            "meta": [
                ("Aplikasi", "Rent Car Platform"),
                ("Audiens", "Mahasiswa, dosen penguji, dan reviewer proyek"),
                ("Fokus", "Alur bisnis rental kendaraan, shuttle, pembayaran, dispatch, dan laporan"),
                ("Tech stack", "Laravel 13, Inertia.js v3, React 19, Tailwind CSS v4"),
            ],
            "abstract": (
                "Dokumen ini merangkum cara menjelaskan Rent Car Platform secara runtut saat presentasi. "
                "Bahasanya dibuat lebih manusiawi, dengan konteks bisnis, alur demo, peran pengguna, "
                "business rules, dan panduan menjawab pertanyaan umum."
            ),
        },
        {
            "type": "section",
            "title": "1. Gambaran Umum Aplikasi",
            "body": [
                (
                    "Rent Car Platform adalah aplikasi web untuk mengelola layanan rental kendaraan dan "
                    "antar-jemput. Sistem ini membantu proses yang biasanya manual - melihat kendaraan, "
                    "membuat pesanan, membayar, memverifikasi pembayaran, menugaskan kendaraan dan supir, "
                    "mencatat pengembalian, sampai membuat laporan."
                ),
                (
                    "Dalam presentasi, posisikan aplikasi ini sebagai solusi operasional untuk usaha rental. "
                    "Customer mendapatkan pengalaman pemesanan yang jelas, sedangkan admin dan kasir memiliki "
                    "alat untuk mengontrol pembayaran, ketersediaan armada, dan status perjalanan."
                ),
            ],
            "callout": {
                "title": "Narasi singkat untuk pembuka",
                "text": (
                    "Aplikasi ini dibuat untuk menyatukan proses customer, admin, kasir, dan supir dalam satu "
                    "alur kerja. Tujuannya bukan hanya menampilkan katalog mobil, tetapi memastikan order bisa "
                    "dibayar, diverifikasi, dijalankan, dikembalikan, dan dilaporkan secara tertib."
                ),
            },
            "bullets": [
                "Rental kendaraan dengan pilihan durasi per jam, hari, minggu, atau bulan.",
                "Layanan shuttle point-to-point berdasarkan tarif rute.",
                "Pembayaran tunai dan transfer manual dengan bukti pembayaran.",
                "Role-based dashboard untuk admin, kasir, customer, dan supir.",
                "Audit trail pada operasi penting seperti pembayaran, dispatch, dan pengembalian.",
            ],
        },
        {
            "type": "section",
            "title": "2. Masalah yang Diselesaikan",
            "body": [
                (
                    "Pada bisnis rental kendaraan, tantangan utama biasanya muncul dari koordinasi manual: "
                    "stok kendaraan tidak sinkron, bukti transfer tercecer, status order tidak jelas, dan laporan "
                    "harus dihitung ulang. Platform ini mengurangi risiko tersebut dengan membuat status dan aksi "
                    "utama berada di dalam sistem."
                )
            ],
            "table": {
                "headers": ["Masalah Operasional", "Dampak", "Solusi di Aplikasi"],
                "rows": [
                    [
                        "Customer sulit melihat kendaraan yang tersedia",
                        "Booking lambat dan sering perlu konfirmasi manual",
                        "Catalog kendaraan, filter, pagination, dan detail kendaraan",
                    ],
                    [
                        "Pembayaran transfer harus dicek satu per satu",
                        "Risiko salah validasi atau bukti hilang",
                        "Upload bukti transfer dan workflow approve/reject",
                    ],
                    [
                        "Admin harus mengingat status kendaraan dan supir",
                        "Kendaraan bisa dipakai tanpa status sistem yang benar",
                        "Dispatch mengubah vehicle menjadi in_use dan driver menjadi on_duty",
                    ],
                    [
                        "Pengembalian terlambat sering dihitung manual",
                        "Overtime tidak konsisten",
                        "Sistem menghitung overtime berdasarkan kelipatan jam",
                    ],
                    [
                        "Laporan pendapatan dibuat manual",
                        "Sulit mengevaluasi performa bisnis",
                        "Report dengan filter tanggal dan revenue paid-only",
                    ],
                ],
            },
        },
        {
            "type": "section",
            "title": "3. Role dan Hak Akses",
            "body": [
                (
                    "Sistem memakai pemisahan role agar setiap pengguna hanya melihat fitur yang sesuai tugasnya. "
                    "Ini penting untuk keamanan, kerapian workflow, dan kemudahan demo karena alur setiap role "
                    "bisa dijelaskan secara terpisah."
                )
            ],
            "table": {
                "headers": ["Role", "Tanggung Jawab Utama", "Dashboard"],
                "rows": [
                    ["Admin", "Mengelola master data, pricing, verifikasi, dispatch, return, dan laporan", "/admin/dashboard"],
                    ["Kasir", "Mencatat pembayaran tunai, memverifikasi transfer, dan mencetak kwitansi", "/kasir/dashboard"],
                    ["Customer", "Melihat katalog, membuat booking, upload bukti transfer, dan melihat receipt", "/catalog"],
                    ["Supir", "Melihat order yang ditugaskan dan menjalankan trip sesuai instruksi", "/driver/dashboard"],
                ],
            },
            "callout": {
                "title": "Poin penting saat menjelaskan role",
                "text": (
                    "Tekankan bahwa redirect setelah login mengikuti role. Ini menunjukkan aplikasi sudah "
                    "memikirkan pengalaman pengguna dan pembatasan akses, bukan sekadar halaman yang berbeda."
                ),
            },
        },
        {
            "type": "section",
            "title": "4. Alur Utama Customer",
            "body": [
                (
                    "Customer adalah titik awal transaksi. Alurnya harus dijelaskan sebagai perjalanan sederhana: "
                    "registrasi, memilih kendaraan, mengisi detail sewa, meninjau harga, membayar, lalu menerima "
                    "kwitansi setelah pembayaran valid."
                )
            ],
            "steps": [
                "Customer membuka halaman register atau login.",
                "Customer masuk ke catalog dan melihat kendaraan dengan pagination empat item per halaman.",
                "Customer memakai filter catalog untuk menyaring kategori atau harga.",
                "Customer memilih kendaraan, menentukan unit sewa, durasi, waktu mulai, dan opsi pickup.",
                "Jika rental keluar kota, sistem menambahkan surcharge sesuai business rule.",
                "Sistem menampilkan ringkasan harga sebelum customer mengonfirmasi order.",
                "Customer memilih pembayaran tunai atau transfer.",
                "Jika transfer, customer upload bukti dan menunggu verifikasi.",
                "Setelah payment paid, customer dapat melihat dan mencetak receipt.",
            ],
            "callout": {
                "title": "Tips demo customer",
                "text": (
                    "Mulailah dari catalog agar audiens langsung melihat nilai aplikasi. Setelah itu baru masuk "
                    "ke detail order dan pembayaran. Hindari terlalu lama di form; cukup jelaskan field yang "
                    "memiliki aturan bisnis seperti minimum tiga jam dan surcharge luar kota."
                ),
            },
        },
        {
            "type": "section",
            "title": "5. Alur Admin dari Master Data sampai Dispatch",
            "body": [
                (
                    "Admin adalah pusat kendali operasional. Admin menyiapkan data kendaraan, kategori, supir, "
                    "aturan harga, tarif shuttle, lalu memproses order yang sudah dibayar."
                )
            ],
            "steps": [
                "Admin mengelola vehicle category, termasuk class level untuk membedakan kelas kendaraan.",
                "Admin menambah dan memperbarui data kendaraan: plate number, brand, model, year, status, dan lokasi.",
                "Admin membuat akun driver dan menjaga status driver agar sesuai kondisi lapangan.",
                "Admin mengatur pricing rule berdasarkan kategori, unit rental, durasi minimum, durasi maksimum, dan rate.",
                "Admin memverifikasi transfer yang masuk dengan approve atau reject disertai alasan.",
                "Admin hanya bisa dispatch order jika payment sudah paid dan order ready_to_dispatch.",
                "Saat dispatch, sistem mengubah order menjadi ongoing, vehicle menjadi in_use, dan driver menjadi on_duty.",
                "Saat kendaraan kembali, admin mencatat actual return time untuk menghitung overtime jika terlambat.",
            ],
            "table": {
                "headers": ["Modul Admin", "Fungsi Presentasi", "Contoh yang Mudah Dijelaskan"],
                "rows": [
                    ["Dashboard", "Menunjukkan kondisi operasional", "Order hari ini, pending payment, armada available"],
                    ["Kategori", "Mengelompokkan kendaraan", "Sedan, MPV, Premium, atau kelas lain"],
                    ["Kendaraan", "Mengelola unit fisik", "Plat nomor unik dan status kendaraan"],
                    ["Harga & Tarif", "Menentukan aturan biaya", "Daily rate, hourly minimum, surcharge"],
                    ["Verifikasi", "Mengamankan pembayaran transfer", "Approve bukti valid, reject bukti salah"],
                    ["Order", "Menjalankan transaksi", "Dispatch dan record return"],
                    ["Laporan", "Membaca performa bisnis", "Revenue paid-only dan transaksi selesai"],
                ],
            },
        },
        {
            "type": "section",
            "title": "6. Alur Kasir",
            "body": [
                (
                    "Kasir menangani sisi pembayaran. Role ini sengaja dibuat lebih terbatas dari admin agar proses "
                    "keuangan bisa dijalankan tanpa memberi akses penuh ke semua master data."
                )
            ],
            "steps": [
                "Kasir mencari order berdasarkan order ID atau nama customer.",
                "Untuk pembayaran tunai, kasir memastikan jumlah sesuai total order.",
                "Setelah dikonfirmasi, payment berubah menjadi paid dan order menjadi ready_to_dispatch.",
                "Sistem membuat receipt dan mencatat audit log.",
                "Kasir juga dapat memeriksa bukti transfer lalu approve atau reject sesuai kewenangan.",
            ],
            "callout": {
                "title": "Poin pembeda kasir dan admin",
                "text": (
                    "Admin mengatur operasional penuh, sedangkan kasir berfokus pada pembayaran. Ini contoh "
                    "penerapan separation of concerns pada level pengguna aplikasi."
                ),
            },
        },
        {
            "type": "section",
            "title": "7. Alur Supir",
            "body": [
                (
                    "Supir menggunakan dashboard untuk melihat order yang ditugaskan. Untuk versi saat ini, "
                    "notifikasi otomatis belum menjadi fokus, sehingga koordinasi awal masih dapat dilakukan "
                    "oleh admin melalui telepon atau WhatsApp."
                )
            ],
            "steps": [
                "Supir login menggunakan akun yang dibuat admin.",
                "Supir melihat detail order yang di-assign: customer, kendaraan, pickup, tujuan, dan jadwal.",
                "Supir mengambil kendaraan dari garasi sesuai koordinasi admin.",
                "Supir menjalankan trip dan mengembalikan kendaraan setelah selesai.",
                "Admin mencatat return agar status order, kendaraan, dan supir kembali sinkron.",
            ],
        },
        {
            "type": "section",
            "title": "8. Shuttle Service",
            "body": [
                (
                    "Selain rental kendaraan, aplikasi menyediakan layanan antar-jemput atau shuttle. Shuttle "
                    "menggunakan pola tarif point-to-point, sehingga harga ditentukan dari area asal dan tujuan "
                    "yang cocok dengan data Shuttle Tariff."
                )
            ],
            "steps": [
                "Customer membuka menu Shuttle Service.",
                "Customer mengisi pickup address, destination address, dan scheduled date time.",
                "Sistem mencari tarif berdasarkan area from dan area to.",
                "Sistem menampilkan estimasi jarak, durasi, dan harga tetap.",
                "Customer mengonfirmasi order dan melanjutkan ke pembayaran.",
            ],
            "callout": {
                "title": "Cara menjelaskan shuttle",
                "text": (
                    "Bedakan shuttle dari rental. Rental bergantung pada kendaraan, durasi, dan opsi pickup. "
                    "Shuttle lebih mirip layanan perjalanan satu rute dengan harga yang sudah ditentukan."
                ),
            },
        },
        {
            "type": "section",
            "title": "9. Business Rules Penting",
            "body": [
                (
                    "Business rules adalah bagian yang biasanya menarik bagi dosen penguji karena menunjukkan "
                    "bahwa aplikasi memiliki logika bisnis, bukan hanya CRUD. Gunakan tabel berikut sebagai "
                    "peta cepat saat presentasi."
                )
            ],
            "table": {
                "headers": ["Kode", "Aturan", "Penjelasan Singkat"],
                "rows": [
                    ["BR-001", "Minimum tiga jam untuk hourly rental", "Jika rental unit hour, durasi minimal adalah 3 jam."],
                    ["BR-002", "Duration-based pricing", "PricingRule dipilih berdasarkan kategori kendaraan, unit rental, dan range durasi."],
                    ["BR-003", "Out-of-town surcharge", "Jika keluar kota, total dikenakan surcharge default 20 persen."],
                    ["BR-004", "Overtime kelipatan jam", "Keterlambatan dihitung ceil(minutes_late / 60) dikali hourly rate."],
                    ["BR-005", "Loyal customer", "Customer dengan minimal satu order selesai bisa memilih driver favorit."],
                    ["BR-006", "Payment lock", "Dispatch hanya boleh dilakukan saat payment paid dan order ready_to_dispatch."],
                    ["BR-007", "Receipt uniqueness", "Nomor kwitansi dibuat unik oleh generator receipt."],
                    ["BR-008", "Vehicle status sync", "Dispatch dan complete otomatis menyinkronkan status kendaraan dan driver."],
                    ["BR-009", "Audit log coverage", "Operasi penting seperti payment, dispatch, dan return tercatat."],
                ],
            },
        },
        {
            "type": "section",
            "title": "10. Skenario Demo yang Disarankan",
            "body": [
                (
                    "Agar presentasi terasa runtut, demo sebaiknya mengikuti satu transaksi dari awal sampai akhir. "
                    "Gunakan data seeded agar waktu demo tidak habis untuk membuat data master dari nol."
                )
            ],
            "steps": [
                "Login sebagai customer dan buka catalog.",
                "Filter kendaraan, pilih satu kendaraan, lalu buat rental order.",
                "Tampilkan review harga dan jelaskan minimum durasi atau surcharge jika dipakai.",
                "Simulasikan pembayaran transfer dengan upload bukti.",
                "Login sebagai admin atau kasir untuk approve pembayaran.",
                "Login sebagai admin, dispatch order yang sudah ready.",
                "Record return dan tunjukkan perubahan status kendaraan serta driver.",
                "Buka receipt atau laporan untuk menutup cerita transaksi.",
            ],
            "table": {
                "headers": ["Menit", "Bagian Demo", "Tujuan"],
                "rows": [
                    ["0-1", "Pembuka", "Jelaskan masalah bisnis dan role pengguna."],
                    ["1-4", "Customer booking", "Tunjukkan nilai utama dari catalog sampai review order."],
                    ["4-6", "Pembayaran", "Jelaskan transfer proof dan approval workflow."],
                    ["6-8", "Admin dispatch", "Tunjukkan validasi paid dan perubahan status operasional."],
                    ["8-9", "Return dan overtime", "Jelaskan logika pengembalian dan biaya terlambat."],
                    ["9-10", "Laporan dan penutup", "Tutup dengan manfaat sistem dan roadmap."],
                ],
            },
        },
        {
            "type": "section",
            "title": "11. Panduan Teknis untuk Developer",
            "body": [
                (
                    "Bagian teknis cukup dijelaskan ringkas saat presentasi, kecuali audiens meminta detail. "
                    "Fokuskan pada arsitektur Route -> Middleware -> Controller -> Service -> Model, karena pola "
                    "ini mudah dipahami dan menunjukkan pemisahan tanggung jawab."
                )
            ],
            "table": {
                "headers": ["Area", "Isi", "Catatan Presentasi"],
                "rows": [
                    ["Backend", "Laravel 13, Fortify, Spatie Permission", "Menangani auth, role, controller, service, dan model."],
                    ["Frontend", "Inertia.js v3, React 19, TypeScript, Tailwind v4", "SPA-like experience tanpa memisahkan API penuh."],
                    ["Testing", "Pest", "Dipakai untuk memastikan pricing, order, dan workflow penting."],
                    ["Tooling", "Pint, Vite, Laravel Brain", "Menjaga style, build frontend, dan dokumentasi arsitektur."],
                ],
            },
            "code": [
                "composer install",
                "npm install",
                "cp .env.example .env",
                "php artisan key:generate",
                "php artisan migrate:fresh --seed",
                "composer run dev",
            ],
        },
        {
            "type": "section",
            "title": "12. Akun Demo Seeded",
            "body": [
                (
                    "Gunakan akun ini untuk mempercepat demo. Sebaiknya siapkan browser atau tab terpisah untuk "
                    "role customer, admin, kasir, dan supir agar perpindahan alur terlihat jelas."
                )
            ],
            "table": {
                "headers": ["Role", "Email", "Password"],
                "rows": [
                    ["Admin", "admin@rentcar.test", "password"],
                    ["Kasir", "kasir@rentcar.test", "password"],
                    ["Customer", "customer@rentcar.test", "password"],
                    ["Driver", "driver@rentcar.test", "password"],
                ],
            },
        },
        {
            "type": "section",
            "title": "13. Troubleshooting Saat Demo",
            "body": [
                (
                    "Bagian ini membantu presenter tetap tenang saat terjadi error umum. Jelaskan penyebabnya "
                    "sebagai validasi sistem, bukan sebagai kegagalan demo."
                )
            ],
            "table": {
                "headers": ["Masalah", "Penyebab", "Solusi Cepat"],
                "rows": [
                    ["Minimum 3 jam", "Rental unit hour dengan durasi kurang dari 3", "Naikkan durasi menjadi 3 jam atau pilih unit day."],
                    ["Cannot dispatch", "Payment belum paid atau order belum ready_to_dispatch", "Approve pembayaran atau input cash terlebih dahulu."],
                    ["Upload bukti gagal", "File lebih dari 5 MB atau format tidak didukung", "Gunakan JPG, PNG, atau PDF di bawah 5 MB."],
                    ["No pricing rule found", "Kategori, unit, dan durasi tidak cocok dengan PricingRule", "Pilih kombinasi lain atau tambahkan rule sebagai admin."],
                    ["Driver belum tahu order", "Notifikasi otomatis belum tersedia", "Admin menghubungi driver secara manual untuk versi ini."],
                ],
            },
        },
        {
            "type": "section",
            "title": "14. FAQ untuk Sesi Tanya Jawab",
            "body": [
                (
                    "Gunakan jawaban berikut sebagai pegangan. Jawaban dibuat singkat agar presenter bisa menjawab "
                    "dengan percaya diri tanpa membuka terlalu banyak detail teknis."
                )
            ],
            "qa": [
                ("Berapa lama verifikasi transfer?", "Manual oleh admin atau kasir. Dalam skenario bisnis, targetnya kurang dari satu jam kerja."),
                ("Apakah customer bisa membatalkan order?", "Untuk versi ini belum menjadi workflow utama. Pembatalan dan refund bisa menjadi roadmap berikutnya."),
                ("Apakah sistem sudah punya payment gateway?", "Belum. Sistem mendukung cash dan transfer manual; virtual account atau payment gateway cocok untuk v2.0."),
                ("Bagaimana jika kendaraan terlambat kembali?", "Admin mencatat actual return time, lalu sistem menghitung overtime berdasarkan kelipatan jam."),
                ("Mengapa loyal customer bisa memilih driver?", "Ini fitur customer retention. Pelanggan yang pernah menyelesaikan order diberi kontrol lebih terhadap pengalaman layanan."),
                ("Apakah aplikasi mobile tersedia?", "Saat ini berbasis web responsive. Aplikasi mobile dapat menjadi pengembangan berikutnya."),
            ],
        },
        {
            "type": "section",
            "title": "15. Roadmap dan Batasan Versi",
            "body": [
                (
                    "Setiap aplikasi punya batasan versi. Menyampaikan batasan dengan jujur justru membuat "
                    "presentasi lebih profesional, karena menunjukkan tim memahami prioritas dan pengembangan bertahap."
                )
            ],
            "table": {
                "headers": ["Area", "Kondisi Saat Ini", "Roadmap"],
                "rows": [
                    ["Notifikasi driver", "Masih koordinasi manual", "Notifikasi real-time atau WhatsApp integration."],
                    ["Cancel dan refund", "Belum menjadi flow utama", "Workflow cancel, refund approval, dan policy biaya."],
                    ["Payment gateway", "Cash dan transfer manual", "Virtual account atau payment gateway otomatis."],
                    ["Availability check", "Perlu penguatan untuk range tanggal", "Pencegahan double-booking berbasis jadwal."],
                    ["Receipt", "Cetak dari browser", "Download PDF receipt langsung dari aplikasi."],
                    ["Corporate account", "Enum customer_type sudah mendukung", "Dashboard khusus B2B dan kontrak pelanggan korporat."],
                ],
            },
        },
        {
            "type": "section",
            "title": "16. Penutup Presentasi",
            "body": [
                (
                    "Tutup presentasi dengan mengulang nilai utama aplikasi: sistem membantu bisnis rental "
                    "mengelola transaksi dari awal sampai akhir secara lebih rapi, terukur, dan mudah diaudit."
                )
            ],
            "callout": {
                "title": "Kalimat penutup yang bisa dipakai",
                "text": (
                    "Dengan Rent Car Platform, proses rental tidak berhenti di katalog kendaraan. Aplikasi ini "
                    "menghubungkan customer, kasir, admin, dan supir dalam satu workflow: booking, pembayaran, "
                    "verifikasi, dispatch, return, receipt, dan laporan."
                ),
            },
        },
    ]


def ensure_dirs() -> None:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    OUT_PDF.parent.mkdir(parents=True, exist_ok=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D7DEE8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_docx_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr[i], "EAF2F8")
        set_cell_border(hdr[i])
        for paragraph in hdr[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(26, 54, 93)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cells[i])
            for paragraph in cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(value) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"
    if widths:
        set_table_width(table, widths)
    doc.add_paragraph()


def add_docx_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    set_cell_border(cell, "B8C7D9")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(20, 83, 45)
    run.font.size = Pt(10)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(2)
    for run in p2.runs:
        run.font.size = Pt(9.5)
    doc.add_paragraph()


def add_docx_bullets(doc: Document, items: list[str], numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build_docx(content: list[dict]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(26, 54, 93)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["List Bullet"].font.name = "Arial"
    styles["List Number"].font.name = "Arial"

    header = section.header.paragraphs[0]
    header.text = "Rent Car Platform - Guide Book Presentasi"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)

    footer = section.footer.paragraphs[0]
    footer.text = "Dokumen pendukung presentasi mahasiswa"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)

    front = content[0]
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(front["title"])
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(front["subtitle"])
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(71, 85, 105)
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rv = version.add_run(VERSION)
    rv.font.size = Pt(10)
    rv.font.color.rgb = RGBColor(100, 116, 139)
    doc.add_paragraph()

    add_docx_table(doc, ["Metadata", "Keterangan"], front["meta"], [1.6, 5.4])
    add_docx_callout(doc, "Ringkasan dokumen", front["abstract"])

    toc_heading = doc.add_heading("Daftar Isi Ringkas", level=1)
    toc_heading.paragraph_format.space_before = Pt(12)
    add_docx_bullets(doc, [item["title"] for item in content[1:]], numbered=True)
    doc.add_page_break()

    for item in content[1:]:
        doc.add_heading(item["title"], level=1)
        for paragraph in item.get("body", []):
            doc.add_paragraph(paragraph)
        if "callout" in item:
            add_docx_callout(doc, item["callout"]["title"], item["callout"]["text"])
        if "bullets" in item:
            add_docx_bullets(doc, item["bullets"])
        if "steps" in item:
            add_docx_bullets(doc, item["steps"], numbered=True)
        if "table" in item:
            headers = item["table"]["headers"]
            widths = None
            if len(headers) == 3:
                widths = [1.25, 2.4, 3.35]
            add_docx_table(doc, headers, item["table"]["rows"], widths)
        if "qa" in item:
            for question, answer in item["qa"]:
                q = doc.add_paragraph()
                qr = q.add_run(f"Q: {question}")
                qr.bold = True
                q.paragraph_format.space_after = Pt(2)
                doc.add_paragraph(f"A: {answer}")
        if "code" in item:
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            set_cell_shading(cell, "0F172A")
            set_cell_border(cell, "0F172A")
            cell.text = "\n".join(item["code"])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(226, 232, 240)
            doc.add_paragraph()
        if item["title"] in {"8. Shuttle Service", "12. Akun Demo Seeded"}:
            doc.add_page_break()

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(OUT_DOCX)


def paragraph_style(name: str, size: int, leading: int, color=colors.HexColor("#111827"), **kwargs) -> ParagraphStyle:
    return ParagraphStyle(
        name,
        fontName="Helvetica",
        fontSize=size,
        leading=leading,
        textColor=color,
        spaceAfter=kwargs.pop("space_after", 8),
        spaceBefore=kwargs.pop("space_before", 0),
        **kwargs,
    )


def pdf_header(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#64748B"))
    canvas.drawString(1.7 * cm, 28.4 * cm, "Rent Car Platform - Guide Book Presentasi")
    canvas.drawRightString(19.3 * cm, 1.2 * cm, f"Halaman {doc.page}")
    canvas.setStrokeColor(colors.HexColor("#CBD5E1"))
    canvas.line(1.7 * cm, 28.15 * cm, 19.3 * cm, 28.15 * cm)
    canvas.restoreState()


def build_pdf(content: list[dict]) -> None:
    doc = BaseDocTemplate(
        str(OUT_PDF),
        pagesize=A4,
        rightMargin=1.7 * cm,
        leftMargin=1.7 * cm,
        topMargin=2.1 * cm,
        bottomMargin=1.8 * cm,
        title=TITLE,
        author="Development Team",
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="content", frames=frame, onPage=pdf_header)])

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CoverTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=26,
        leading=32,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1A365D"),
        spaceAfter=14,
    )
    subtitle_style = paragraph_style("Subtitle", 12, 16, colors.HexColor("#475569"), alignment=TA_CENTER, space_after=18)
    h1 = paragraph_style("H1", 15, 19, colors.HexColor("#1A365D"), space_before=14, space_after=8)
    h1.fontName = "Helvetica-Bold"
    body = paragraph_style("Body", 9.5, 13, colors.HexColor("#111827"), space_after=7)
    callout_title = paragraph_style("CalloutTitle", 9.5, 12, colors.HexColor("#14532D"), space_after=4)
    callout_title.fontName = "Helvetica-Bold"
    code_style = paragraph_style("Code", 8.5, 11, colors.HexColor("#E2E8F0"), space_after=0)
    code_style.fontName = "Courier"

    story = []
    front = content[0]
    story.append(Spacer(1, 2.2 * cm))
    story.append(Paragraph(front["title"], title_style))
    story.append(Paragraph(front["subtitle"], subtitle_style))
    story.append(Paragraph(VERSION, subtitle_style))
    story.append(Spacer(1, 0.7 * cm))
    meta_data = [[Paragraph(f"<b>{a}</b>", body), Paragraph(b, body)] for a, b in front["meta"]]
    meta = Table(meta_data, colWidths=[4.2 * cm, 12.4 * cm])
    meta.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EAF2F8")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(meta)
    story.append(Spacer(1, 0.4 * cm))
    story.append(callout_block(front["abstract"], callout_title, body, "Ringkasan dokumen"))
    story.append(PageBreak())

    story.append(Paragraph("Daftar Isi Ringkas", h1))
    story.append(make_list([item["title"] for item in content[1:]], body, numbered=True))
    story.append(PageBreak())

    for item in content[1:]:
        story.append(Paragraph(item["title"], h1))
        for para in item.get("body", []):
            story.append(Paragraph(para, body))
        if "callout" in item:
            story.append(callout_block(item["callout"]["text"], callout_title, body, item["callout"]["title"]))
        if "bullets" in item:
            story.append(make_list(item["bullets"], body))
        if "steps" in item:
            story.append(make_list(item["steps"], body, numbered=True))
        if "table" in item:
            story.append(pdf_table(item["table"]["headers"], item["table"]["rows"], body))
            story.append(Spacer(1, 0.2 * cm))
        if "qa" in item:
            for question, answer in item["qa"]:
                story.append(Paragraph(f"<b>Q: {question}</b>", body))
                story.append(Paragraph(f"A: {answer}", body))
        if "code" in item:
            rows = [[Paragraph(line, code_style)] for line in item["code"]]
            code_table = Table(rows, colWidths=[16.6 * cm])
            code_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#0F172A")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#0F172A")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(code_table)
            story.append(Spacer(1, 0.2 * cm))
        if item["title"] in {"8. Shuttle Service", "12. Akun Demo Seeded"}:
            story.append(PageBreak())

    doc.build(story)


def callout_block(text: str, title_style: ParagraphStyle, body_style: ParagraphStyle, title: str):
    block = Table(
        [[Paragraph(title, title_style)], [Paragraph(text, body_style)]],
        colWidths=[16.6 * cm],
    )
    block.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#B8C7D9")),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    return KeepTogether([block, Spacer(1, 0.25 * cm)])


def make_list(items: list[str], style: ParagraphStyle, numbered: bool = False):
    flowables = []
    for idx, item in enumerate(items, 1):
        bullet = f"{idx}." if numbered else "-"
        flowables.append(ListItem(Paragraph(item, style), bulletText=bullet, leftIndent=12))
    return ListFlowable(flowables, bulletType="1" if numbered else "bullet", start="1", leftIndent=16)


def pdf_table(headers: list[str], rows: list[list[str]], style: ParagraphStyle):
    data = [[Paragraph(f"<b>{h}</b>", style) for h in headers]]
    for row in rows:
        data.append([Paragraph(str(cell), style) for cell in row])
    if len(headers) == 2:
        col_widths = [4.4 * cm, 12.2 * cm]
    elif len(headers) == 3:
        col_widths = [3.3 * cm, 5.0 * cm, 8.3 * cm]
    else:
        col_widths = [16.6 * cm / len(headers)] * len(headers)
    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF2F8")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1A365D")),
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return table


def build_markdown(content: list[dict]) -> None:
    lines = [
        f"# {TITLE}",
        "",
        f"> {SUBTITLE}",
        f"> {VERSION}",
        "",
        "Dokumen ini adalah versi presentasi yang sudah dirapikan dari `docs/GUIDE_BOOK.md`.",
        "Bahasa dibuat lebih natural untuk mahasiswa, dengan alur demo, business rules, FAQ, dan roadmap.",
        "",
        "## Metadata",
        "",
    ]
    for key, value in content[0]["meta"]:
        lines.append(f"- **{key}:** {value}")
    lines.extend(["", f"> **Ringkasan:** {content[0]['abstract']}", ""])
    for item in content[1:]:
        lines.extend([f"## {item['title']}", ""])
        for para in item.get("body", []):
            lines.extend([para, ""])
        if "callout" in item:
            lines.extend([f"> **{item['callout']['title']}**", f"> {item['callout']['text']}", ""])
        if "bullets" in item:
            for bullet in item["bullets"]:
                lines.append(f"- {bullet}")
            lines.append("")
        if "steps" in item:
            for idx, step in enumerate(item["steps"], 1):
                lines.append(f"{idx}. {step}")
            lines.append("")
        if "table" in item:
            headers = item["table"]["headers"]
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in item["table"]["rows"]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
        if "qa" in item:
            for question, answer in item["qa"]:
                lines.extend([f"**Q: {question}**  ", f"A: {answer}", ""])
        if "code" in item:
            lines.extend(["```bash", *item["code"], "```", ""])
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    content = build_content()
    build_markdown(content)
    build_docx(content)
    build_pdf(content)
    print(OUT_MD)
    print(OUT_DOCX)
    print(OUT_PDF)


if __name__ == "__main__":
    main()
